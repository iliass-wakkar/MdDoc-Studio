"""
Native Standalone Python Markdown-to-DOCX Converter.
Zero binary dependencies — pure Python using python-docx, markdown-it-py / markdown, and PyYAML.
"""

import os
import re
from datetime import datetime
from typing import Optional, Dict, Any, Tuple, List

import yaml
from bs4 import BeautifulSoup, NavigableString, Tag
import markdown
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

from .themes import get_theme, THEMES
from .oxml import (
    hex_to_rgb,
    set_cell_shading,
    set_cell_margins,
    set_cell_borders,
    set_row_cant_split,
    set_row_header,
    set_p_spacing,
    set_p_shading,
    add_p_border_bottom,
    add_p_border_left,
    add_toc_field,
    add_page_number_fields,
    add_horizontal_rule,
)


def parse_frontmatter(md_content: str) -> Tuple[Dict[str, Any], str]:
    """Extract YAML frontmatter and body markdown from markdown text."""
    frontmatter = {}
    body = md_content

    pattern = r'^---\s*\n(.*?)\n---\s*\n'
    match = re.match(pattern, md_content, re.DOTALL)
    if match:
        fm_text = match.group(1)
        try:
            frontmatter = yaml.safe_load(fm_text) or {}
        except Exception:
            frontmatter = {}
        body = md_content[match.end():]

    return frontmatter, body


class MarkdownToDocxConverter:
    """
    High-fidelity Markdown to DOCX renderer inspired by Kimi Skills design heuristics.
    """
    def __init__(
        self,
        theme_name: str = "modern",
        custom_theme: Optional[Dict[str, Any]] = None,
        title: Optional[str] = None,
        subtitle: Optional[str] = None,
        author: Optional[str] = None,
        date: Optional[str] = None,
        show_cover: bool = True,
        show_toc: bool = True,
        page_size: str = "A4",
    ):
        if custom_theme and isinstance(custom_theme, dict):
            self.theme = custom_theme
            self.theme_name = "custom"
        else:
            self.theme_name = theme_name
            self.theme = get_theme(theme_name)

        self.title = title
        self.subtitle = subtitle
        self.author = author
        self.date = date or datetime.now().strftime("%B %d, %Y")
        self.show_cover = show_cover
        self.show_toc = show_toc
        self.page_size = page_size
        self.heading_count = 0
        self.doc = docx.Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        """Configure default styles, margins, and headers/footers."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.theme["font_body"]
        font.size = Pt(10.5)
        font.color.rgb = hex_to_rgb(self.theme["text"])
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

        for section in self.doc.sections:
            if self.page_size.upper() == "LETTER":
                section.page_width = Inches(8.5)
                section.page_height = Inches(11.0)
            else:  # Default A4
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)

            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.different_first_page_header_footer = True

    def _render_cover_page(self):
        """Render a polished, executive cover page."""
        # Top spacing
        for _ in range(3):
            p = self.doc.add_paragraph()
            set_p_spacing(p, after=Pt(12))

        # Top Accent Decorative Line
        p_acc = self.doc.add_paragraph()
        p_acc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_acc = p_acc.add_run("━" * 18)
        run_acc.font.size = Pt(14)
        run_acc.font.color.rgb = hex_to_rgb(self.theme["secondary"])
        set_p_spacing(p_acc, after=Pt(36))

        # Document Title
        if self.title:
            p_title = self.doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(self.title)
            run_title.font.size = Pt(32)
            run_title.font.bold = True
            run_title.font.color.rgb = hex_to_rgb(self.theme["primary"])
            run_title.font.name = self.theme["font_heading"]
            set_p_spacing(p_title, after=Pt(12))

        # Subtitle
        if self.subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(self.subtitle)
            run_sub.font.size = Pt(15)
            run_sub.font.color.rgb = hex_to_rgb(self.theme["light_text"])
            run_sub.font.name = self.theme["font_body"]
            set_p_spacing(p_sub, after=Pt(36))

        # Divider Bar
        p_div = self.doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_div = p_div.add_run("─" * 28)
        run_div.font.size = Pt(10)
        run_div.font.color.rgb = hex_to_rgb(self.theme["secondary"])
        set_p_spacing(p_div, after=Pt(40))

        # Author
        if self.author:
            p_auth = self.doc.add_paragraph()
            p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_auth = p_auth.add_run(f"Author: {self.author}")
            run_auth.font.size = Pt(11.5)
            run_auth.font.color.rgb = hex_to_rgb(self.theme["text"])
            run_auth.font.name = self.theme["font_body"]
            run_auth.font.italic = True
            set_p_spacing(p_auth, after=Pt(6))

        # Date
        if self.date:
            p_date = self.doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_date = p_date.add_run(self.date)
            run_date.font.size = Pt(10)
            run_date.font.color.rgb = hex_to_rgb(self.theme["light_text"])
            run_date.font.name = self.theme["font_body"]
            set_p_spacing(p_date, after=Pt(36))

        # Bottom Accent
        for _ in range(2):
            p = self.doc.add_paragraph()
            set_p_spacing(p, after=Pt(12))

        p_bot = self.doc.add_paragraph()
        p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bot = p_bot.add_run("━" * 18)
        run_bot.font.size = Pt(14)
        run_bot.font.color.rgb = hex_to_rgb(self.theme["secondary"])

        self.doc.add_page_break()

    def _render_headers_footers(self):
        """Setup headers and footers for body pages."""
        for section in self.doc.sections:
            # Header
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if self.title:
                run_h = hp.add_run(self.title)
                run_h.font.size = Pt(8.5)
                run_h.font.color.rgb = hex_to_rgb(self.theme["light_text"])
                run_h.font.name = self.theme["font_body"]
            set_p_spacing(hp, before=Pt(0), after=Pt(4))
            add_p_border_bottom(hp, self.theme["table_border"], sz=4, space_after=Pt(4))

            # Footer
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_p_spacing(fp, before=Pt(4), after=Pt(0))
            add_page_number_fields(
                fp,
                prefix="Page ",
                separator=" of ",
                text_color=self.theme["light_text"],
                font_name=self.theme["font_body"],
                font_size=Pt(9)
            )

    def _process_inline(self, paragraph, element, is_bold=False, is_italic=False, is_code=False):
        """Recursively process inline markdown elements (bold, italic, code, links)."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    if is_code:
                        run.font.name = self.theme["font_code"]
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = hex_to_rgb(self.theme["accent"])
                    else:
                        run.font.name = self.theme["font_body"]
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = hex_to_rgb(self.theme["text"])
                    if is_bold:
                        run.font.bold = True
                    if is_italic:
                        run.font.italic = True
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                if tag_name in ('strong', 'b'):
                    self._process_inline(paragraph, child, is_bold=True, is_italic=is_italic, is_code=is_code)
                elif tag_name in ('em', 'i'):
                    self._process_inline(paragraph, child, is_bold=is_bold, is_italic=True, is_code=is_code)
                elif tag_name == 'code':
                    self._process_inline(paragraph, child, is_bold=is_bold, is_italic=is_italic, is_code=True)
                elif tag_name == 'a':
                    link_text = child.get_text()
                    run = paragraph.add_run(link_text)
                    run.font.name = self.theme["font_body"]
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = hex_to_rgb(self.theme["link"])
                    run.font.underline = True
                    if is_bold:
                        run.font.bold = True
                    if is_italic:
                        run.font.italic = True
                elif tag_name == 'del':
                    run = paragraph.add_run(child.get_text())
                    run.font.strike = True
                else:
                    self._process_inline(paragraph, child, is_bold=is_bold, is_italic=is_italic, is_code=is_code)

    def _add_heading(self, text: str, level: int):
        """Add styled heading with keep-with-next and hierarchy colors."""
        p = self.doc.add_paragraph()
        if level == 1:
            if self.heading_count > 0:
                self.doc.add_page_break()
            run = p.add_run(text)
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(self.theme["heading1"])
            run.font.name = self.theme["font_heading"]
            set_p_spacing(p, before=Pt(20), after=Pt(8), keep_with_next=True)
            add_p_border_bottom(p, self.theme["secondary"], sz=16, space_after=Pt(12))
        elif level == 2:
            run = p.add_run(text)
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(self.theme["heading2"])
            run.font.name = self.theme["font_heading"]
            set_p_spacing(p, before=Pt(16), after=Pt(6), keep_with_next=True)
        elif level == 3:
            run = p.add_run(text)
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(self.theme["heading3"])
            run.font.name = self.theme["font_heading"]
            set_p_spacing(p, before=Pt(12), after=Pt(4), keep_with_next=True)
        else:  # H4, H5, H6
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(self.theme["heading4"])
            run.font.name = self.theme["font_heading"]
            set_p_spacing(p, before=Pt(8), after=Pt(2), keep_with_next=True)

        self.heading_count += 1

    def _add_paragraph(self, element: Tag):
        """Add styled standard paragraph."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._process_inline(p, element)
        set_p_spacing(p, before=Pt(2), after=Pt(6), line_spacing=1.15)

    def _add_blockquote_or_alert(self, element: Tag):
        """
        Add blockquote or GitHub-style Admonition ([!NOTE], [!TIP], [!WARNING], etc.).
        """
        raw_text = element.get_text().strip()
        alert_match = re.match(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*(.*)', raw_text, re.IGNORECASE | re.DOTALL)

        if alert_match:
            alert_type = alert_match.group(1).upper()
            alerts = self.theme.get("alerts", {})
            alert_cfg = alerts.get(alert_type, alerts.get("NOTE", {}))
            
            # Single-cell callout container table for crisp borders & padding
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Cm(16.0)

            cell = table.cell(0, 0)
            set_cell_shading(cell, alert_cfg.get("bg", "F0F9FF"))
            set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=10, right_pt=10)
            set_cell_borders(
                cell,
                left={'val': 'single', 'sz': 24, 'color': alert_cfg.get("color", self.theme["secondary"])},
                top={'val': 'none'},
                bottom={'val': 'none'},
                right={'val': 'none'}
            )

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_p_spacing(p, before=Pt(2), after=Pt(4))

            # Header badge
            icon = alert_cfg.get("icon", "")
            title = alert_cfg.get("title", alert_type)
            run_badge = p.add_run(f"{icon} {title}\n")
            run_badge.font.bold = True
            run_badge.font.size = Pt(10.5)
            run_badge.font.color.rgb = hex_to_rgb(alert_cfg.get("color", self.theme["secondary"]))
            run_badge.font.name = self.theme["font_heading"]

            # Content paragraphs inside blockquote
            p_children = element.find_all('p')
            if p_children:
                for idx, cp in enumerate(p_children):
                    if idx == 0:
                        clean_html = re.sub(r'^\s*\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*', '', str(cp), flags=re.IGNORECASE)
                        temp_soup = BeautifulSoup(clean_html, 'html.parser')
                        self._process_inline(p, temp_soup)
                    else:
                        np = cell.add_paragraph()
                        set_p_spacing(np, before=Pt(2), after=Pt(4))
                        self._process_inline(np, cp)
            else:
                run_text = p.add_run(alert_match.group(2))
                run_text.font.size = Pt(10)
                run_text.font.color.rgb = hex_to_rgb(self.theme["text"])
                run_text.font.name = self.theme["font_body"]

            spacer = self.doc.add_paragraph()
            set_p_spacing(spacer, after=Pt(6))
        else:
            # Standard Blockquote
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Cm(16.0)

            cell = table.cell(0, 0)
            set_cell_shading(cell, self.theme["quote_bg"])
            set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=10, right_pt=10)
            set_cell_borders(
                cell,
                left={'val': 'single', 'sz': 24, 'color': self.theme["quote_border"]},
                top={'val': 'none'},
                bottom={'val': 'none'},
                right={'val': 'none'}
            )

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_p_spacing(p, before=Pt(2), after=Pt(2))

            for child_p in element.find_all('p') or [element]:
                self._process_inline(p, child_p, is_italic=True)

            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = hex_to_rgb(self.theme["quote_text"])
                run.font.name = self.theme["font_heading"]

            spacer = self.doc.add_paragraph()
            set_p_spacing(spacer, after=Pt(6))

    def _add_code_block(self, code_text: str):
        """Add styled code block with background and border."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(16.0)

        cell = table.cell(0, 0)
        set_cell_shading(cell, self.theme["code_bg"])
        set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
        border_col = self.theme["code_border"]
        set_cell_borders(
            cell,
            top={'val': 'single', 'sz': 4, 'color': border_col},
            bottom={'val': 'single', 'sz': 4, 'color': border_col},
            left={'val': 'single', 'sz': 4, 'color': border_col},
            right={'val': 'single', 'sz': 4, 'color': border_col}
        )

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_p_spacing(p, before=Pt(2), after=Pt(2), line_spacing=1.0)

        lines = code_text.rstrip('\n').split('\n')
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.font.name = self.theme["font_code"]
            run.font.size = Pt(9)
            run.font.color.rgb = hex_to_rgb(self.theme["code_text"])
            if i < len(lines) - 1:
                run_br = p.add_run('\n')
                run_br.font.name = self.theme["font_code"]

        spacer = self.doc.add_paragraph()
        set_p_spacing(spacer, after=Pt(6))

    def _add_list(self, element: Tag, ordered: bool = False):
        """Add styled bullet, numbered, or task list."""
        items = element.find_all('li', recursive=False)
        for idx, li in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_p_spacing(p, before=Pt(1), after=Pt(3))
            p.paragraph_format.left_indent = Inches(0.25)

            li_text = li.get_text()
            # Task list detection: [ ] or [x]
            if li_text.startswith('[ ] ') or li_text.startswith('[x] ') or li_text.startswith('[X] '):
                is_checked = not li_text.startswith('[ ] ')
                icon = "☑ " if is_checked else "☐ "
                run_icon = p.add_run(icon)
                run_icon.font.name = "Segoe UI Symbol"
                run_icon.font.bold = True
                run_icon.font.color.rgb = hex_to_rgb(self.theme["secondary"] if is_checked else self.theme["light_text"])
                
                # strip prefix from first child
                first_str = next(li.strings, '')
                cleaned_str = re.sub(r'^\[[ xX]\]\s*', '', first_str)
                run_text = p.add_run(cleaned_str)
                run_text.font.name = self.theme["font_body"]
                run_text.font.size = Pt(10.5)
                run_text.font.color.rgb = hex_to_rgb(self.theme["text"])
            else:
                if ordered:
                    run_bullet = p.add_run(f"{idx}. ")
                else:
                    run_bullet = p.add_run("•  ")

                run_bullet.font.bold = True
                run_bullet.font.color.rgb = hex_to_rgb(self.theme["secondary"])
                run_bullet.font.name = self.theme["font_body"]
                self._process_inline(p, li)

    def _add_table(self, element: Tag):
        """Add professional booktabs / zebra styled table."""
        rows = element.find_all('tr')
        if not rows:
            return

        cols_count = max(len(r.find_all(['td', 'th'])) for r in rows)
        if cols_count == 0:
            return

        table = self.doc.add_table(rows=len(rows), cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for i, row in enumerate(rows):
            set_row_cant_split(table.rows[i])
            is_header = (i == 0 or row.parent.name == 'thead')
            if is_header:
                set_row_header(table.rows[i])

            cells = row.find_all(['td', 'th'])
            for j, cell_elem in enumerate(cells):
                if j >= cols_count:
                    break
                docx_cell = table.cell(i, j)
                docx_cell.text = ""
                set_cell_margins(docx_cell, top_pt=5, bottom_pt=5, left_pt=8, right_pt=8)

                border_color = self.theme["table_border"]
                set_cell_borders(
                    docx_cell,
                    top={'val': 'single', 'sz': 4, 'color': border_color},
                    bottom={'val': 'single', 'sz': 6 if is_header else 4, 'color': border_color},
                    left={'val': 'single', 'sz': 4, 'color': border_color},
                    right={'val': 'single', 'sz': 4, 'color': border_color}
                )

                p = docx_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_p_spacing(p, before=Pt(2), after=Pt(2))
                self._process_inline(p, cell_elem)

                if is_header:
                    set_cell_shading(docx_cell, self.theme["table_header_bg"])
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = hex_to_rgb(self.theme["table_header_text"])
                        run.font.name = self.theme["font_body"]
                else:
                    if i % 2 == 0 and self.theme.get("table_row_alt"):
                        set_cell_shading(docx_cell, self.theme["table_row_alt"])
                    for run in p.runs:
                        run.font.color.rgb = hex_to_rgb(self.theme["text"])
                        run.font.name = self.theme["font_body"]

        spacer = self.doc.add_paragraph()
        set_p_spacing(spacer, after=Pt(6))

    def _add_image(self, src: str, alt: str = ""):
        """Add centered image with optional italic caption."""
        try:
            if os.path.exists(src):
                self.doc.add_picture(src, width=Inches(5.8))
                last_p = self.doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_p_spacing(last_p, before=Pt(8), after=Pt(4))

                if alt:
                    p_cap = self.doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_cap = p_cap.add_run(f"Figure: {alt}")
                    run_cap.font.size = Pt(9)
                    run_cap.font.italic = True
                    run_cap.font.color.rgb = hex_to_rgb(self.theme["light_text"])
                    set_p_spacing(p_cap, before=Pt(2), after=Pt(8))
            else:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"📷 [Image: {alt or src}]")
                run.font.size = Pt(9.5)
                run.font.italic = True
                run.font.color.rgb = hex_to_rgb(self.theme["light_text"])
                set_p_spacing(p, before=Pt(6), after=Pt(6))
        except Exception as e:
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Image load error: {str(e)}]")
            run.font.color.rgb = hex_to_rgb(self.theme["accent"])

    def _process_element(self, element: Tag):
        """Dispatch processing for each HTML tag."""
        if isinstance(element, NavigableString):
            return

        tag = element.name.lower()
        if tag == 'h1':
            self._add_heading(element.get_text(strip=True), 1)
        elif tag == 'h2':
            self._add_heading(element.get_text(strip=True), 2)
        elif tag == 'h3':
            self._add_heading(element.get_text(strip=True), 3)
        elif tag in ('h4', 'h5', 'h6'):
            self._add_heading(element.get_text(strip=True), 4)
        elif tag == 'p':
            # Check if paragraph only contains an image
            img = element.find('img', recursive=False)
            if img and len(element.contents) == 1:
                self._add_image(img.get('src', ''), img.get('alt', ''))
            else:
                self._add_paragraph(element)
        elif tag == 'blockquote':
            self._add_blockquote_or_alert(element)
        elif tag == 'pre':
            code = element.find('code')
            code_text = code.get_text() if code else element.get_text()
            self._add_code_block(code_text)
        elif tag == 'ul':
            self._add_list(element, ordered=False)
        elif tag == 'ol':
            self._add_list(element, ordered=True)
        elif tag == 'table':
            self._add_table(element)
        elif tag == 'img':
            self._add_image(element.get('src', ''), element.get('alt', ''))
        elif tag == 'hr':
            add_horizontal_rule(self.doc, self.theme["hr_color"])
        elif tag in ('div', 'section', 'article', 'main'):
            for child in element.children:
                if isinstance(child, Tag):
                    self._process_element(child)

    def convert(self, md_content: str) -> docx.Document:
        """Convert Markdown string into styled Document object."""
        frontmatter, body_md = parse_frontmatter(md_content)

        # Merge frontmatter metadata
        if "title" in frontmatter and not self.title:
            self.title = str(frontmatter["title"])
        if "subtitle" in frontmatter and not self.subtitle:
            self.subtitle = str(frontmatter["subtitle"])
        if "author" in frontmatter and not self.author:
            self.author = str(frontmatter["author"])
        if "date" in frontmatter and not self.date:
            self.date = str(frontmatter["date"])
        if "theme" in frontmatter:
            self.theme_name = str(frontmatter["theme"])
            self.theme = get_theme(self.theme_name)
        if "cover_page" in frontmatter:
            self.show_cover = bool(frontmatter["cover_page"])
        if "toc" in frontmatter:
            self.show_toc = bool(frontmatter["toc"])

        # Fallback title from first H1 if still missing
        if not self.title:
            h1_match = re.search(r'^#\s+(.+)$', body_md, re.MULTILINE)
            if h1_match:
                self.title = h1_match.group(1).strip()

        # 1. Cover Page
        if self.show_cover:
            self._render_cover_page()

        # 2. Table of Contents
        if self.show_toc:
            add_toc_field(self.doc, self.theme)
            self.doc.add_page_break()

        # 3. Render Body Elements
        html = markdown.markdown(
            body_md,
            extensions=['tables', 'fenced_code', 'nl2br', 'sane_lists']
        )
        soup = BeautifulSoup(html, 'html.parser')
        root = soup.body if soup.body else soup

        for elem in root.children:
            if isinstance(elem, Tag):
                self._process_element(elem)

        # 4. Headers and Footers
        self._render_headers_footers()

        return self.doc

    def save(self, output_path: str):
        """Save generated DOCX file."""
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        self.doc.save(output_path)


def convert_markdown_to_docx(
    input_path: str,
    output_path: Optional[str] = None,
    theme_name: str = "modern",
    title: Optional[str] = None,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    date: Optional[str] = None,
    show_cover: bool = True,
    show_toc: bool = True,
    page_size: str = "A4"
) -> str:
    """Convenience helper function to convert markdown file to docx."""
    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    if output_path is None:
        base = os.path.splitext(input_path)[0]
        output_path = base + ".docx"

    converter = MarkdownToDocxConverter(
        theme_name=theme_name,
        title=title,
        subtitle=subtitle,
        author=author,
        date=date,
        show_cover=show_cover,
        show_toc=show_toc,
        page_size=page_size
    )

    converter.convert(md_text)
    converter.save(output_path)
    return os.path.abspath(output_path)
