"""
Low-level OpenXML (OXML) helpers for python-docx.
Enables precise control over Word styling, shading, borders, margins, and fields.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches, Cm


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex string (e.g. '1E3A5F' or '#1E3A5F') to RGBColor."""
    clean_hex = hex_color.lstrip('#')
    if len(clean_hex) == 3:
        clean_hex = ''.join(c * 2 for c in clean_hex)
    return RGBColor(
        int(clean_hex[0:2], 16),
        int(clean_hex[2:4], 16),
        int(clean_hex[4:6], 16)
    )


def set_cell_shading(cell, fill_hex: str):
    """Apply background color shading to a table cell."""
    clean_hex = fill_hex.lstrip('#')
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), clean_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top_pt=5, bottom_pt=5, left_pt=8, right_pt=8):
    """
    Set inner cell padding (margins) in points.
    Converted to dxa (1 pt = 20 dxa).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val_pt in [('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)]:
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(int(val_pt * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, **kwargs):
    """
    Set custom borders on a cell.
    kwargs: top, bottom, left, right dicts with keys:
    - val: 'single', 'nil', 'dashed', etc.
    - sz: size in 1/8 pt (e.g. 4 = 0.5pt, 8 = 1pt, 24 = 3pt)
    - color: hex color
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', 'auto').lstrip('#'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_row_cant_split(row):
    """Ensure table row does not split awkwardly across page breaks."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def set_row_header(row):
    """Mark table row as a repeating header on subsequent pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))


def set_p_spacing(paragraph, before=Pt(0), after=Pt(0), line_spacing=1.15, keep_with_next=False):
    """Configure paragraph spacing and keep-with-next rules."""
    p_format = paragraph.paragraph_format
    p_format.space_before = before
    p_format.space_after = after
    p_format.line_spacing = line_spacing
    if keep_with_next:
        p_format.keep_with_next = True


def set_p_shading(paragraph, fill_hex: str):
    """Set background color shading on an entire paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    clean_hex = fill_hex.lstrip('#')
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), clean_hex)
    pPr.append(shading)


def add_p_border_bottom(paragraph, color_hex: str, sz: int = 12, space_after=Pt(10)):
    """Add a colored bottom border line beneath a paragraph (e.g. for H1)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)
    paragraph.paragraph_format.space_after = space_after


def add_p_border_left(paragraph, color_hex: str, sz: int = 24):
    """Add a thick colored left border accent to a paragraph (for blockquotes/callouts)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))  # 24 = 3pt
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(left)
    pPr.append(pBdr)


def add_field_code(paragraph, field_type: str, placeholder: str = "1"):
    """Insert a dynamic Word field code (e.g. PAGE, NUMPAGES, TOC)."""
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_type} '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run(placeholder)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


def add_page_number_fields(paragraph, prefix="Page ", separator=" of ", text_color=None, font_name='Calibri', font_size=Pt(9)):
    """Add 'Page X of Y' field codes to a footer paragraph."""
    run_prefix = paragraph.add_run(prefix)
    run_prefix.font.size = font_size
    run_prefix.font.name = font_name
    if text_color:
        run_prefix.font.color.rgb = hex_to_rgb(text_color)

    # Current Page
    run_p1 = paragraph.add_run()
    f1 = OxmlElement('w:fldChar')
    f1.set(qn('w:fldCharType'), 'begin')
    run_p1._r.append(f1)
    
    run_p2 = paragraph.add_run()
    instr1 = OxmlElement('w:instrText')
    instr1.set(qn('xml:space'), 'preserve')
    instr1.text = ' PAGE '
    run_p2._r.append(instr1)
    
    run_p3 = paragraph.add_run()
    f2 = OxmlElement('w:fldChar')
    f2.set(qn('w:fldCharType'), 'separate')
    run_p3._r.append(f2)
    
    run_p4 = paragraph.add_run("1")
    run_p4.font.size = font_size
    run_p4.font.name = font_name
    if text_color:
        run_p4.font.color.rgb = hex_to_rgb(text_color)
        
    run_p5 = paragraph.add_run()
    f3 = OxmlElement('w:fldChar')
    f3.set(qn('w:fldCharType'), 'end')
    run_p5._r.append(f3)

    # Separator
    run_sep = paragraph.add_run(separator)
    run_sep.font.size = font_size
    run_sep.font.name = font_name
    if text_color:
        run_sep.font.color.rgb = hex_to_rgb(text_color)

    # Total Pages
    run_n1 = paragraph.add_run()
    nf1 = OxmlElement('w:fldChar')
    nf1.set(qn('w:fldCharType'), 'begin')
    run_n1._r.append(nf1)
    
    run_n2 = paragraph.add_run()
    ninstr = OxmlElement('w:instrText')
    ninstr.set(qn('xml:space'), 'preserve')
    ninstr.text = ' NUMPAGES '
    run_n2._r.append(ninstr)
    
    run_n3 = paragraph.add_run()
    nf2 = OxmlElement('w:fldChar')
    nf2.set(qn('w:fldCharType'), 'separate')
    run_n3._r.append(nf2)
    
    run_n4 = paragraph.add_run("1")
    run_n4.font.size = font_size
    run_n4.font.name = font_name
    if text_color:
        run_n4.font.color.rgb = hex_to_rgb(text_color)
        
    run_n5 = paragraph.add_run()
    nf3 = OxmlElement('w:fldChar')
    nf3.set(qn('w:fldCharType'), 'end')
    run_n5._r.append(nf3)


def add_toc_field(doc, theme, title="Table of Contents"):
    """Insert a styled Table of Contents container with Word field code."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_t = p_title.add_run(title)
    run_t.font.size = Pt(20)
    run_t.font.bold = True
    run_t.font.color.rgb = hex_to_rgb(theme["heading1"])
    run_t.font.name = theme["font_heading"]
    set_p_spacing(p_title, before=Pt(18), after=Pt(8), keep_with_next=True)
    add_p_border_bottom(p_title, theme["secondary"], sz=12, space_after=Pt(14))

    p_toc = doc.add_paragraph()
    add_field_code(p_toc, 'TOC \\o "1-3" \\h \\z \\u', placeholder="Table of contents entries will display here.")
    set_p_spacing(p_toc, before=Pt(6), after=Pt(6))

    p_hint = doc.add_paragraph("(Right-click TOC in Microsoft Word and select 'Update Field' to refresh page numbers)")
    p_hint.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_hint = p_hint.runs[0]
    run_hint.font.size = Pt(8.5)
    run_hint.font.color.rgb = hex_to_rgb(theme["light_text"])
    run_hint.font.italic = True
    set_p_spacing(p_hint, before=Pt(4), after=Pt(18))
    return p_toc


def add_horizontal_rule(doc, color_hex: str):
    """Add a sleek divider line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p_border_bottom(p, color_hex, sz=6, space_after=Pt(12))
    set_p_spacing(p, before=Pt(12), after=Pt(12))
    return p
