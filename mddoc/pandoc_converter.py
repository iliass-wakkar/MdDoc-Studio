"""
Pandoc conversion engine with automated template selection and post-processing polish.
"""

import os
import re
import subprocess
from typing import Optional, Dict, Any

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from .themes import get_theme
from .templates import find_pandoc_executable, generate_reference_docx
from .oxml import (
    hex_to_rgb,
    set_cell_shading,
    set_cell_margins,
    set_cell_borders,
    set_row_cant_split,
    set_row_header,
    set_p_spacing,
    add_p_border_left,
    set_p_shading,
)


def is_pandoc_available() -> bool:
    """Check if Pandoc binary is found on the system."""
    return find_pandoc_executable() is not None


def _post_process_docx(docx_path: str, theme: Dict[str, Any]):
    """
    Post-process Pandoc output DOCX to enhance tables, callouts, and layout.
    """
    doc = docx.Document(docx_path)

    # 1. Enhance all Tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = table.rows
        if not rows:
            continue

        for i, row in enumerate(rows):
            set_row_cant_split(row)
            is_header = (i == 0)
            if is_header:
                set_row_header(row)

            for j, cell in enumerate(row.cells):
                set_cell_margins(cell, top_pt=5, bottom_pt=5, left_pt=8, right_pt=8)
                
                # Borders
                border_color = theme["table_border"]
                set_cell_borders(
                    cell,
                    top={'val': 'single', 'sz': 4, 'color': border_color},
                    bottom={'val': 'single', 'sz': 6 if is_header else 4, 'color': border_color},
                    left={'val': 'single', 'sz': 4, 'color': border_color},
                    right={'val': 'single', 'sz': 4, 'color': border_color}
                )

                if is_header:
                    set_cell_shading(cell, theme["table_header_bg"])
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        set_p_spacing(p, before=Pt(2), after=Pt(2))
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = hex_to_rgb(theme["table_header_text"])
                            run.font.name = theme["font_body"]
                else:
                    if i % 2 == 0 and theme.get("table_row_alt"):
                        set_cell_shading(cell, theme["table_row_alt"])
                    for p in cell.paragraphs:
                        set_p_spacing(p, before=Pt(2), after=Pt(2))
                        for run in p.runs:
                            run.font.color.rgb = hex_to_rgb(theme["text"])
                            run.font.name = theme["font_body"]

    # 2. Check for GFM Admonitions in paragraphs (e.g. [!NOTE], [!WARNING])
    alerts = theme.get("alerts", {})
    for p in doc.paragraphs:
        p_text = p.text.strip()
        match = re.match(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*(.*)', p_text, re.IGNORECASE)
        if match:
            alert_type = match.group(1).upper()
            rest_text = match.group(2)
            alert_cfg = alerts.get(alert_type, alerts.get("NOTE", {}))
            
            p.text = ""  # clear and rebuild
            add_p_border_left(p, alert_cfg["color"], sz=24)
            set_p_shading(p, alert_cfg["bg"])
            set_p_spacing(p, before=Pt(6), after=Pt(6))
            p.paragraph_format.left_indent = Inches(0.3)
            
            # Badge run
            icon = alert_cfg.get("icon", "")
            title = alert_cfg.get("title", alert_type)
            run_badge = p.add_run(f"{icon} {title}: ")
            run_badge.font.bold = True
            run_badge.font.color.rgb = hex_to_rgb(alert_cfg["color"])
            run_badge.font.name = theme["font_body"]
            
            if rest_text:
                run_body = p.add_run(rest_text)
                run_body.font.color.rgb = hex_to_rgb(theme["text"])
                run_body.font.name = theme["font_body"]

    doc.save(docx_path)


def convert_with_pandoc(
    input_path: str,
    output_path: str,
    theme_name: str = "modern",
    toc: bool = True,
    reference_docx: Optional[str] = None,
    extra_args: Optional[list] = None
) -> str:
    """
    Convert Markdown to DOCX using Pandoc + Reference Template + Post-processor.
    """
    pandoc_exe = find_pandoc_executable()
    if not pandoc_exe:
        raise RuntimeError("Pandoc executable not found on system. Please install Pandoc or use the native engine.")

    theme = get_theme(theme_name)

    # Use or generate reference docx
    if not reference_docx or not os.path.isfile(reference_docx):
        reference_docx = generate_reference_docx(theme_name)

    cmd = [
        pandoc_exe,
        input_path,
        "-o", output_path,
        f"--reference-doc={reference_docx}",
    ]

    if toc:
        cmd.append("--toc")

    if extra_args:
        cmd.extend(extra_args)

    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"Pandoc conversion failed:\n{res.stderr}")

    # Post-process styling
    _post_process_docx(output_path, theme)

    return os.path.abspath(output_path)
