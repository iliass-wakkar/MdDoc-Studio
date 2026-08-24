"""
Reference DOCX template generator for Pandoc and standalone Word workflows.
Generates styled reference docx templates for each color theme.
"""

import os
import shutil
import subprocess
from typing import Dict, Any, Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

from .themes import THEMES, get_theme
from .oxml import (
    hex_to_rgb,
    set_cell_shading,
    set_cell_margins,
    set_cell_borders,
    add_p_border_bottom,
    add_p_border_left,
    set_p_shading,
    set_p_spacing,
    add_page_number_fields,
)


def find_pandoc_executable() -> Optional[str]:
    """Find pandoc executable on PATH or standard install locations."""
    # Check PATH first
    cmd = shutil.which("pandoc")
    if cmd:
        return cmd
    
    # Check standard Windows paths
    candidates = [
        os.path.expandvars(r"%LOCALAPPDATA%\Pandoc\pandoc.exe"),
        os.path.expandvars(r"%APPDATA%\Pandoc\pandoc.exe"),
        r"C:\Program Files\Pandoc\pandoc.exe",
        r"C:\Program Files (x86)\Pandoc\pandoc.exe",
        r"C:\Program Files\RStudio\resources\app\bin\quarto\bin\tools\pandoc.exe",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def get_or_create_base_reference_docx(cache_path: Optional[str] = None) -> str:
    """
    Extract default reference.docx from pandoc if available,
    or generate a clean base docx if pandoc is absent.
    """
    if cache_path and os.path.exists(cache_path):
        return cache_path

    out_file = cache_path or os.path.join(os.path.dirname(__file__), "base_reference.docx")
    pandoc_exe = find_pandoc_executable()

    if pandoc_exe:
        try:
            res = subprocess.run([pandoc_exe, "--print-default-data-file", "reference.docx"],
                                 capture_output=True, check=True)
            with open(out_file, "wb") as f:
                f.write(res.stdout)
            return out_file
        except Exception:
            pass

    # Fallback: create base docx with python-docx
    doc = docx.Document()
    doc.save(out_file)
    return out_file


def _apply_theme_to_styles(doc: docx.Document, theme: Dict[str, Any]):
    """Apply theme typography, colors, borders, and shading to docx styles."""
    styles = doc.styles

    # 1. Normal / Body Text
    for s_name in ['Normal', 'Body Text', 'Compact']:
        if s_name in styles:
            s = styles[s_name]
            s.font.name = theme["font_body"]
            s.font.size = Pt(10.5)
            s.font.color.rgb = hex_to_rgb(theme["text"])
            s.paragraph_format.line_spacing = 1.15
            s.paragraph_format.space_after = Pt(4)

    # 2. Title
    if 'Title' in styles:
        s = styles['Title']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(28)
        s.font.bold = True
        s.font.color.rgb = hex_to_rgb(theme["primary"])
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(36)
        s.paragraph_format.space_after = Pt(12)

    # 3. Subtitle
    if 'Subtitle' in styles:
        s = styles['Subtitle']
        s.font.name = theme["font_body"]
        s.font.size = Pt(14)
        s.font.color.rgb = hex_to_rgb(theme["light_text"])
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(24)

    # 4. Author
    if 'Author' in styles:
        s = styles['Author']
        s.font.name = theme["font_body"]
        s.font.size = Pt(11)
        s.font.italic = True
        s.font.color.rgb = hex_to_rgb(theme["text"])
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(6)

    # 5. Date
    if 'Date' in styles:
        s = styles['Date']
        s.font.name = theme["font_body"]
        s.font.size = Pt(10)
        s.font.color.rgb = hex_to_rgb(theme["light_text"])
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(24)

    # 6. Heading 1
    if 'Heading 1' in styles:
        s = styles['Heading 1']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(20)
        s.font.bold = True
        s.font.color.rgb = hex_to_rgb(theme["heading1"])
        s.paragraph_format.space_before = Pt(24)
        s.paragraph_format.space_after = Pt(10)
        s.paragraph_format.keep_with_next = True
        # Bottom border on H1 style
        pPr = s._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '16')  # 2pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), theme["secondary"].lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)

    # 7. Heading 2
    if 'Heading 2' in styles:
        s = styles['Heading 2']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(15)
        s.font.bold = True
        s.font.color.rgb = hex_to_rgb(theme["heading2"])
        s.paragraph_format.space_before = Pt(16)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    # 8. Heading 3
    if 'Heading 3' in styles:
        s = styles['Heading 3']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(12.5)
        s.font.bold = True
        s.font.color.rgb = hex_to_rgb(theme["heading3"])
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.keep_with_next = True

    # 9. Heading 4
    if 'Heading 4' in styles:
        s = styles['Heading 4']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(11)
        s.font.bold = True
        s.font.color.rgb = hex_to_rgb(theme["heading4"])
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(2)
        s.paragraph_format.keep_with_next = True

    # 10. Block Text (Blockquotes)
    if 'Block Text' in styles:
        s = styles['Block Text']
        s.font.name = theme["font_heading"]
        s.font.size = Pt(10.5)
        s.font.italic = True
        s.font.color.rgb = hex_to_rgb(theme["quote_text"])
        s.paragraph_format.left_indent = Inches(0.4)
        s.paragraph_format.right_indent = Inches(0.2)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.15
        
        # Left border + background shading
        pPr = s._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')  # 3pt
        left.set(qn('w:space'), '12')
        left.set(qn('w:color'), theme["quote_border"].lstrip('#'))
        pBdr.append(left)
        pPr.append(pBdr)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), theme["quote_bg"].lstrip('#'))
        pPr.append(shd)

    # 11. Source Code (Code blocks)
    for code_style_name in ['Source Code', 'Verbatim Char']:
        if code_style_name in styles:
            s = styles[code_style_name]
            s.font.name = theme["font_code"]
            s.font.size = Pt(9)
            s.font.color.rgb = hex_to_rgb(theme["code_text"])
            if s.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
                s.paragraph_format.left_indent = Inches(0.2)
                s.paragraph_format.space_before = Pt(4)
                s.paragraph_format.space_after = Pt(4)
                s.paragraph_format.line_spacing = 1.0
                pPr = s._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), theme["code_bg"].lstrip('#'))
                pPr.append(shd)

    # 12. Hyperlink
    if 'Hyperlink' in styles:
        s = styles['Hyperlink']
        s.font.color.rgb = hex_to_rgb(theme["link"])
        s.font.underline = True


def _configure_document_layout(doc: docx.Document, theme: Dict[str, Any]):
    """Set margins, different first page, header and footer on sections."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.different_first_page_header_footer = True

        # Header (pages 2+)
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_p_spacing(hp, before=Pt(0), after=Pt(4))
        # Subtle bottom border line on header
        add_p_border_bottom(hp, theme["table_border"], sz=4, space_after=Pt(4))

        # Footer (pages 2+)
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_p_spacing(fp, before=Pt(4), after=Pt(0))
        add_page_number_fields(fp, prefix="Page ", separator=" of ",
                               text_color=theme["light_text"],
                               font_name=theme["font_body"],
                               font_size=Pt(9))


def generate_reference_docx(theme_name: str = "modern", output_path: Optional[str] = None) -> str:
    """
    Generate a styled reference.docx for the specified theme.
    Returns the absolute path to the generated template file.
    """
    theme = get_theme(theme_name)
    base_docx = get_or_create_base_reference_docx()
    doc = docx.Document(base_docx)

    _apply_theme_to_styles(doc, theme)
    _configure_document_layout(doc, theme)

    if output_path is None:
        templates_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
        os.makedirs(templates_dir, exist_ok=True)
        output_path = os.path.join(templates_dir, f"reference-{theme_name.lower()}.docx")
    
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


def generate_all_reference_templates(output_dir: Optional[str] = None) -> Dict[str, str]:
    """Generate reference.docx files for all available themes."""
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
    os.makedirs(output_dir, exist_ok=True)

    results = {}
    for theme_name in THEMES.keys():
        target_file = os.path.join(output_dir, f"reference-{theme_name}.docx")
        results[theme_name] = generate_reference_docx(theme_name, target_file)
    return results
