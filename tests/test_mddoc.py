"""
Automated Test Suite for MdDoc.
Tests themes, template generation, native converter, pandoc converter, and CLI.
"""

import os
import unittest
import docx
from mddoc.themes import THEMES, get_theme
from mddoc.templates import generate_reference_docx, generate_all_reference_templates
from mddoc.native_converter import convert_markdown_to_docx, parse_frontmatter
from mddoc.pandoc_converter import convert_with_pandoc, is_pandoc_available


class TestMdDoc(unittest.TestCase):
    def setUp(self):
        self.output_dir = os.path.join(os.path.dirname(__file__), "test_outputs")
        os.makedirs(self.output_dir, exist_ok=True)
        self.sample_md = os.path.join(self.output_dir, "test_doc.md")
        with open(self.sample_md, "w", encoding="utf-8") as f:
            f.write("""---
title: "Test Unit Document"
subtitle: "Sub Unit Test"
author: "QA Bot"
date: "August 2026"
theme: "nordic"
toc: true
cover_page: true
---

# Main Section

This is a test paragraph with **bold**, *italic*, and `inline_code`.

> [!NOTE]
> This is a test note alert.

> [!WARNING]
> This is a test warning alert.

```python
def test_func():
    return True
```

| Header A | Header B |
|---|---|
| Value 1 | Value 2 |
| Value 3 | Value 4 |

- Item 1
- Item 2
- [x] Task done
- [ ] Task pending
""")

    def test_themes_completeness(self):
        """Ensure all themes have required keys."""
        required_keys = [
            "name", "font_heading", "font_body", "font_code",
            "primary", "secondary", "accent", "heading1", "heading2",
            "text", "code_bg", "quote_border", "quote_bg",
            "table_header_bg", "table_border"
        ]
        for theme_name, theme_data in THEMES.items():
            for k in required_keys:
                self.assertIn(k, theme_data, f"Theme '{theme_name}' missing key '{k}'")

    def test_frontmatter_parsing(self):
        """Test YAML frontmatter parsing."""
        md_text = "---\ntitle: Hello\nauthor: World\n---\n# Body Content"
        fm, body = parse_frontmatter(md_text)
        self.assertEqual(fm.get("title"), "Hello")
        self.assertEqual(fm.get("author"), "World")
        self.assertEqual(body.strip(), "# Body Content")

    def test_template_generation(self):
        """Test generating reference docx files for all themes."""
        res = generate_all_reference_templates(self.output_dir)
        for theme_name in THEMES.keys():
            t_path = res[theme_name]
            self.assertTrue(os.path.exists(t_path))
            self.assertGreater(os.path.getsize(t_path), 1000)
            # Verify it opens as a valid docx
            doc = docx.Document(t_path)
            self.assertIsNotNone(doc)

    def test_native_converter_all_themes(self):
        """Test native python conversion for all themes."""
        for theme_name in THEMES.keys():
            out_file = os.path.join(self.output_dir, f"native_{theme_name}.docx")
            res_path = convert_markdown_to_docx(
                input_path=self.sample_md,
                output_path=out_file,
                theme_name=theme_name
            )
            self.assertTrue(os.path.exists(res_path))
            self.assertGreater(os.path.getsize(res_path), 5000)
            doc = docx.Document(res_path)
            self.assertGreater(len(doc.paragraphs), 5)

    def test_pandoc_converter(self):
        """Test Pandoc conversion engine if pandoc is available."""
        if not is_pandoc_available():
            self.skipTest("Pandoc not installed on system")
        for theme_name in THEMES.keys():
            out_file = os.path.join(self.output_dir, f"pandoc_{theme_name}.docx")
            res_path = convert_with_pandoc(
                input_path=self.sample_md,
                output_path=out_file,
                theme_name=theme_name,
                toc=True
            )
            self.assertTrue(os.path.exists(res_path))
            self.assertGreater(os.path.getsize(res_path), 5000)
    def test_web_api(self):
        """Test MarkdownToDocxConverter in-memory generation used by Web API."""
        from mddoc.native_converter import MarkdownToDocxConverter
        import io
        converter = MarkdownToDocxConverter(theme_name="corporate", title="Web API Test")
        doc = converter.convert("# Web Heading\n\nWeb body test.")
        bio = io.BytesIO()
        doc.save(bio)
        data = bio.getvalue()
        self.assertGreater(len(data), 2000)
        # Verify valid docx from bytes
        read_doc = docx.Document(io.BytesIO(data))
        self.assertIsNotNone(read_doc)


if __name__ == "__main__":
    unittest.main()
