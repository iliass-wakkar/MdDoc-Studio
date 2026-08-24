"""
Comprehensive Verification Script for MdDoc Desktop GUI & Web Studio.
"""

import io
import os
import sys

# Ensure repository root is on sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import json
import time
import urllib.request
import threading
import tkinter as tk
import docx

# 1. Test Tkinter Desktop GUI Lifecycle
print("=" * 60)
print("TEST 1: Testing Tkinter Desktop GUI Lifecycle...")
print("=" * 60)

from mddoc.gui import MdDocGUI

root = tk.Tk()
# Hide window during automated test
root.withdraw()

app = MdDocGUI(root)
sample_md = os.path.abspath("examples/comprehensive_sample.md")
out_docx = os.path.abspath("examples/gui_test_output.docx")

app.var_input_file.set(sample_md)
app.var_output_file.set(out_docx)
app.var_theme.set("nordic")
app._on_theme_changed()

print("  -> Initialized Tkinter GUI window structure: SUCCESS")
print(f"  -> Selected Theme: {app.var_theme.get()} ({app.lbl_theme_desc.cget('text')})")

# Run the conversion worker directly
print("  -> Executing GUI conversion worker...")
app._run_conversion_worker(sample_md, out_docx)

# Verify output file
if os.path.exists(out_docx) and os.path.getsize(out_docx) > 5000:
    doc = docx.Document(out_docx)
    print(f"  -> Output DOCX created successfully ({os.path.getsize(out_docx)} bytes, {len(doc.paragraphs)} paragraphs)")
    print("  -> Desktop GUI test: [PASSED]")
else:
    print("  -> Desktop GUI test: [FAILED]")
    sys.exit(1)

root.destroy()


# 2. Test Local Web Studio Server Lifecycle
print("\n" + "=" * 60)
print("TEST 2: Testing Embedded Web Studio Server & API...")
print("=" * 60)

from mddoc.web_server import HTTPServer, MdDocRequestHandler, find_free_port

test_port = find_free_port(9876)
server = HTTPServer(('127.0.0.1', test_port), MdDocRequestHandler)
server_thread = threading.Thread(target=server.serve_forever, daemon=True)
server_thread.start()
time.sleep(0.5)

base_url = f"http://127.0.0.1:{test_port}"
print(f"  -> Web Studio server running at {base_url}")

# Test 2a: GET /
req = urllib.request.Request(base_url)
with urllib.request.urlopen(req) as resp:
    status = resp.getcode()
    html = resp.read().decode('utf-8')
    assert status == 200
    assert "MdDoc Studio" in html
    print(f"  -> GET / returned HTTP {status} (HTML length: {len(html)} bytes): [PASSED]")

# Test 2b: POST /api/convert
post_data = {
    "markdown": "# Web API Conversion Test\n\n> [!NOTE]\n> Verified via automated test.\n\n| Column 1 | Column 2 |\n|---|---|\n| A | B |",
    "theme": "academic",
    "title": "Automated Web Test",
    "author": "Test Suite",
    "cover_page": True,
    "toc": True,
    "page_size": "A4"
}
post_bytes = json.dumps(post_data).encode('utf-8')
post_req = urllib.request.Request(
    f"{base_url}/api/convert",
    data=post_bytes,
    headers={"Content-Type": "application/json"},
    method="POST"
)

with urllib.request.urlopen(post_req) as resp:
    status = resp.getcode()
    docx_data = resp.read()
    content_type = resp.headers.get("Content-Type")
    assert status == 200
    assert "openxmlformats" in content_type
    assert len(docx_data) > 3000
    # Validate it's a real Word docx
    test_doc = docx.Document(io.BytesIO(docx_data))
    print(f"  -> POST /api/convert returned HTTP {status} ({len(docx_data)} bytes DOCX payload): [PASSED]")
    print(f"  -> Verified generated DOCX content: '{test_doc.paragraphs[0].text or 'Valid document'}'")

server.shutdown()
print("  -> Web Studio server shut down cleanly.")
print("  -> Web Studio test: [PASSED]")

print("\n" + "=" * 60)
print("ALL GUI & WEB STUDIO TESTS PASSED SUCCESSFULLY! [100% OK]")
print("=" * 60)
